from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

out_path = Path(r"e:/Ajaxx/Projects/ecoguard/AeroGaurd/AeroGuard_Report_Template.docx")

doc = Document()

# Title page
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("AeroGuard Project Report")
run.bold = True
run.font.size = Pt(24)

doc.add_paragraph("")

for line in [
    "Hyper-Local Air Quality and Health Risk Forecasting",
    "",
    "Name: ________________________________",
    "Roll No / ID: _________________________",
    "Department: __________________________",
    "Institute: ____________________________",
    "Date: ________________________________",
]:
    para = doc.add_paragraph(line)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


doc.add_page_break()

# Core sections with empty placeholders
sections = [
    ("1. Abstract", [
        "Write a short summary of your project in 150-250 words.",
        "",
        "[Leave this section text-only or fill later.]",
    ]),
    ("2. Problem Statement", [
        "Describe the real-world problem and motivation.",
        "",
        "[Paste screenshot if needed]",
        "",
    ]),
    ("3. Objectives", [
        "- Objective 1",
        "- Objective 2",
        "- Objective 3",
    ]),
    ("4. Dataset Description", [
        "Source file(s): _________________________",
        "Total rows: _____________________________",
        "Total locations/cities: _________________",
        "Date range: _____________________________",
        "",
        "Screenshot Placeholder: Dataset preview / schema",
        "",
        "[PASTE SCREENSHOT HERE]",
        "",
    ]),
    ("5. System Architecture", [
        "Explain modules and data flow.",
        "",
        "Screenshot Placeholder: Architecture diagram",
        "[PASTE SCREENSHOT HERE]",
        "",
    ]),
    ("6. Project Modules", [
        "6.1 Data Loading and Cleaning",
        "Screenshot Placeholder:",
        "[PASTE SCREENSHOT HERE]",
        "",
        "6.2 Forecasting Model",
        "Screenshot Placeholder:",
        "[PASTE SCREENSHOT HERE]",
        "",
        "6.3 Health Risk and Persona Advice",
        "Screenshot Placeholder:",
        "[PASTE SCREENSHOT HERE]",
        "",
        "6.4 Smart Action Engine",
        "Screenshot Placeholder:",
        "[PASTE SCREENSHOT HERE]",
        "",
        "6.5 Streamlit Dashboard",
        "Screenshot Placeholder:",
        "[PASTE SCREENSHOT HERE]",
    ]),
    ("7. ML Model Evaluation", [
        "Model used: ____________________________",
        "Train/Test split: ______________________",
        "Accuracy: ______________________________",
        "",
        "Screenshot Placeholder: Terminal output (metrics)",
        "[PASTE SCREENSHOT HERE]",
        "",
        "Screenshot Placeholder: Confusion matrix (PNG)",
        "[PASTE SCREENSHOT HERE]",
    ]),
    ("8. Dashboard Screenshots", [
        "8.1 Live Dashboard",
        "[PASTE SCREENSHOT HERE]",
        "",
        "8.2 AI Prediction Page",
        "[PASTE SCREENSHOT HERE]",
        "",
        "8.3 Heatmap and Location Page",
        "[PASTE SCREENSHOT HERE]",
        "",
        "8.4 Smart City Actions",
        "[PASTE SCREENSHOT HERE]",
    ]),
    ("9. Results and Observations", [
        "Write key findings after adding screenshots.",
        "",
        "- Observation 1",
        "- Observation 2",
        "- Observation 3",
    ]),
    ("10. Challenges Faced", [
        "- Challenge 1 and how you handled it",
        "- Challenge 2 and how you handled it",
    ]),
    ("11. Future Scope", [
        "- Future enhancement 1",
        "- Future enhancement 2",
        "- Future enhancement 3",
    ]),
    ("12. Conclusion", [
        "Write final conclusion after compiling the report.",
    ]),
]

for title, lines in sections:
    h = doc.add_heading(title, level=1)
    for line in lines:
        doc.add_paragraph(line)

# Appendix section for extra screenshot pages
doc.add_page_break()
doc.add_heading("Appendix: Additional Screenshots", level=1)
for i in range(1, 7):
    doc.add_paragraph(f"Appendix Screenshot {i}")
    doc.add_paragraph("[PASTE SCREENSHOT HERE]")
    doc.add_paragraph("")

doc.save(out_path)
print(f"Created: {out_path}")
